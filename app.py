import os
import sqlite3
import subprocess
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, jsonify, get_flashed_messages, send_file
from werkzeug.utils import secure_filename
from werkzeug.security import check_password_hash, generate_password_hash
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
import ffmpeg
from datetime import datetime
import fitz  # PyMuPDF para PDFs
from docx import Document  # python-docx para DOC/DOCX
import io
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import uuid
import pyotp
import qrcode
from io import BytesIO
import base64
import logging

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['THUMBNAIL_FOLDER'] = 'static/thumbnails'
app.config['IMAGE_FOLDER'] = 'static/images'
app.config['VIDEO_FOLDER'] = 'static/videos'
app.config['PREVIEW_FOLDER'] = 'static/previews'
app.config['MAX_CONTENT_LENGTH'] = 2 * 1024 * 1024 * 1024  # 2 GB
app.config['FFMPEG_PATH'] = r"C:\Program Files\ffmpeg\bin\ffmpeg.exe"  # Ajuste este caminho conforme necessário
app.secret_key = 'supersecretkey'
app.config['SECRET_KEY'] = 'sua-chave-secreta-aqui'  # Substitua por uma chave forte
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = False  # False para localhost


from flask import Flask, render_template

# ... (outras importações e configurações do app)

# Manipulador de erros para 500
@app.errorhandler(500)
def internal_server_error(e):
    return render_template('offline.html'), 500

# Manipulador genérico para outros erros (opcional)
@app.errorhandler(Exception)
def handle_exception(e):
    logging.error(f"Erro não tratado: {str(e)}")
    return render_template('offline.html'), 500

# Flask-Login setup
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.unauthorized_handler
def unauthorized():
    flash('Por favor, faça login para acessar esta página.', 'warning')
    return redirect(url_for('login'))

# Ensure folders exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['THUMBNAIL_FOLDER'], exist_ok=True)
os.makedirs(app.config['IMAGE_FOLDER'], exist_ok=True)
os.makedirs(app.config['VIDEO_FOLDER'], exist_ok=True)
os.makedirs(app.config['PREVIEW_FOLDER'], exist_ok=True)

# User class for Flask-Login
class User(UserMixin):
    def __init__(self, id, username, email, role, otp_secret=None):
        self.id = id
        self.username = username
        self.email = email
        self.role = role
        self.otp_secret = otp_secret

    @property
    def is_admin(self):
        return self.role == 'admin'

@login_manager.user_loader
def load_user(user_id):
    conn = get_db_connection()
    user = conn.execute('SELECT * FROM users WHERE id = ?', (user_id,)).fetchone()
    conn.close()
    if user is None:
        return None
    otp_secret = user['otp_secret'] if 'otp_secret' in user else None  # Lidar com ausência de otp_secret
    return User(user[0], user[1], user[2], user[4], otp_secret) if not hasattr(user, 'keys') else User(user['id'], user['username'], user['email'], user['role'], otp_secret)

def allowed_file(filename, allowed_extensions=None):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'mp4', 'mov', 'mkv', 'avi', 'png', 'jpg', 'jpeg', 'gif', 'wmv', 'pdf', 'doc', 'docx'}

def get_db_connection():
    conn = sqlite3.connect('database.db', timeout=10)  # Adiciona timeout para evitar bloqueios
    conn.row_factory = sqlite3.Row
    return conn

def generate_preview(file_path, filename):
    ext = filename.rsplit('.', 1)[1].lower()
    preview_path = None
    preview_text = None

    try:
        if ext == 'pdf':
            doc = fitz.open(file_path)
            if doc.page_count > 0:
                page = doc.load_page(0)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                preview_filename = f"{os.path.splitext(filename)[0]}_preview.png"
                preview_path = os.path.join(app.config['PREVIEW_FOLDER'], preview_filename).replace('\\', '/')
                os.makedirs(app.config['PREVIEW_FOLDER'], exist_ok=True)
                pix.save(preview_path)
                preview_path = f'previews/{preview_filename}'
                preview_text = "Preview gerado a partir de PDF (texto não extraível)."  # Fallback
        elif ext in ('doc', 'docx'):
            doc = Document(file_path)
            preview_text = ''
            for para in doc.paragraphs[:3]:
                if para.text.strip():
                    preview_text += para.text + '\n'
            preview_text = preview_text.strip()[:300] + ('...' if len(preview_text) > 300 else '')
            if not preview_text:
                preview_text = "Nenhum texto encontrado no Word."
        return preview_path, preview_text
    except Exception as e:
        print(f"[ERROR] Erro ao gerar preview para {filename}: {str(e)}")
        return None, None

def init_db():
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.executescript('''
        CREATE TABLE IF NOT EXISTS articles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            question TEXT NOT NULL,
            solution TEXT NOT NULL,
            created_at TEXT,
            created_by INTEGER,
            product TEXT,
            subproduct TEXT,
            modified_by_username TEXT,
            modified_at TEXT,
            FOREIGN KEY (created_by) REFERENCES users (id)
        );
        CREATE TABLE IF NOT EXISTS media (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            article_id INTEGER,
            media_path TEXT,
            thumbnail_path TEXT,
            FOREIGN KEY (article_id) REFERENCES articles (id)
        );
        CREATE TABLE IF NOT EXISTS settings (
            id INTEGER PRIMARY KEY,
            logo_path TEXT
        );
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL UNIQUE,
            email TEXT NOT NULL UNIQUE,
            password_hash TEXT NOT NULL,
            role TEXT NOT NULL DEFAULT 'user',
            otp_secret TEXT
        );
        CREATE TABLE IF NOT EXISTS tutorials (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            description TEXT NOT NULL,
            product TEXT NOT NULL,
            subproduct TEXT,
            video_path TEXT,
            thumbnail_path TEXT,
            created_at TEXT,
            created_by INTEGER,
            modified_by_username TEXT,
            modified_at TEXT,
            FOREIGN KEY (created_by) REFERENCES users (id)
        );
        CREATE TABLE IF NOT EXISTS procedures (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            file_path TEXT NOT NULL,
            preview_path TEXT,
            preview_text TEXT,
            created_at TEXT,
            created_by INTEGER,
            modified_by_username TEXT,
            modified_at TEXT,
            product TEXT,
            subproduct TEXT,
            FOREIGN KEY (created_by) REFERENCES users (id)
        );
        INSERT OR IGNORE INTO settings (id, logo_path) VALUES (1, 'images/default-logo.png');
    ''')
    cursor.execute('PRAGMA table_info(articles)')
    columns = [col['name'] for col in cursor.fetchall()]
    if 'created_at' not in columns:
        cursor.execute('ALTER TABLE articles ADD COLUMN created_at TEXT')
    if 'created_by' not in columns:
        cursor.execute('ALTER TABLE articles ADD COLUMN created_by INTEGER')
    if 'product' not in columns:
        cursor.execute('ALTER TABLE articles ADD COLUMN product TEXT')
    if 'subproduct' not in columns:
        cursor.execute('ALTER TABLE articles ADD COLUMN subproduct TEXT')
    if 'modified_by_username' not in columns:
        cursor.execute('ALTER TABLE articles ADD COLUMN modified_by_username TEXT')
    if 'modified_at' not in columns:
        cursor.execute('ALTER TABLE articles ADD COLUMN modified_at TEXT')
    cursor.execute("UPDATE articles SET created_at = ? WHERE created_at IS NULL", (datetime.now().strftime('%Y-%m-%d %H:%M:%S'),))
    cursor.execute("UPDATE articles SET created_by = 1 WHERE created_by IS NULL")
    cursor.execute("UPDATE articles SET product = 'Outros' WHERE product IS NULL")
    cursor.execute("UPDATE articles SET subproduct = '' WHERE subproduct IS NULL")

    cursor.execute('PRAGMA table_info(tutorials)')
    columns = [col['name'] for col in cursor.fetchall()]
    if 'modified_by_username' not in columns:
        cursor.execute('ALTER TABLE tutorials ADD COLUMN modified_by_username TEXT')
    if 'modified_at' not in columns:
        cursor.execute('ALTER TABLE tutorials ADD COLUMN modified_at TEXT')

    cursor.execute('PRAGMA table_info(procedures)')
    columns = [col['name'] for col in cursor.fetchall()]
    if 'created_at' not in columns:
        cursor.execute('ALTER TABLE procedures ADD COLUMN created_at TEXT')
    if 'created_by' not in columns:
        cursor.execute('ALTER TABLE procedures ADD COLUMN created_by INTEGER')
    if 'modified_by_username' not in columns:
        cursor.execute('ALTER TABLE procedures ADD COLUMN modified_by_username TEXT')
    if 'modified_at' not in columns:
        cursor.execute('ALTER TABLE procedures ADD COLUMN modified_at TEXT')
    if 'preview_path' not in columns:
        cursor.execute('ALTER TABLE procedures ADD COLUMN preview_path TEXT')
    if 'preview_text' not in columns:
        cursor.execute('ALTER TABLE procedures ADD COLUMN preview_text TEXT')
    if 'product' not in columns:
        cursor.execute('ALTER TABLE procedures ADD COLUMN product TEXT')
        cursor.execute("UPDATE procedures SET product = 'Outros' WHERE product IS NULL")
    if 'subproduct' not in columns:
        cursor.execute('ALTER TABLE procedures ADD COLUMN subproduct TEXT')
        cursor.execute("UPDATE procedures SET subproduct = '' WHERE subproduct IS NULL")
    cursor.execute("UPDATE procedures SET created_at = ? WHERE created_at IS NULL", (datetime.now().strftime('%Y-%m-%d %H:%M:%S'),))
    cursor.execute("UPDATE procedures SET created_by = 1 WHERE created_by IS NULL")

    cursor.execute("SELECT * FROM users WHERE username = 'admin'")
    if not cursor.fetchone():
        cursor.execute("INSERT INTO users (username, email, password_hash, role) VALUES (?, ?, ?, ?)", 
                       ('admin', 'admin@infoaxis.com.br', generate_password_hash('admin123'), 'admin'))

    # Adicionar coluna otp_secret se não existir
    cursor.execute('PRAGMA table_info(users)')
    columns = [col['name'] for col in cursor.fetchall()]
    if 'otp_secret' not in columns:
        cursor.execute('ALTER TABLE users ADD COLUMN otp_secret TEXT')

    conn.commit()
    conn.close()
    print("Database initialized or tables verified.")

init_db()

from flask import request, url_for
from flask import session
from flask_login import current_user, login_required

@app.route('/')
@app.route('/index')
@login_required
def index():
    search_query = request.args.get('search', '')
    product_filter = request.args.get('product', '')
    subproduct_filter = request.args.get('subproduct', '')
    sort_order = request.args.get('sort', 'desc')

    conn = get_db_connection()
    query = '''
        SELECT a.*, u.username as created_by_username
        FROM articles a
        JOIN users u ON a.created_by = u.id
        WHERE 1=1
    '''
    params = []

    if search_query:
        query += ' AND (a.title LIKE ? OR a.question LIKE ? OR a.solution LIKE ?)'
        params.extend([f'%{search_query}%', f'%{search_query}%', f'%{search_query}%'])

    if product_filter:
        query += ' AND a.product = ?'
        params.append(product_filter)

    if subproduct_filter:
        query += ' AND a.subproduct = ?'
        params.append(subproduct_filter)

    query += ' ORDER BY a.created_at ' + ('DESC' if sort_order == 'desc' else 'ASC')

    articles = conn.execute(query, params).fetchall()
    articles_list = []

    for article in articles:
        media = conn.execute('SELECT * FROM media WHERE article_id = ?', (article['id'],)).fetchall()
        article_dict = dict(article)
        article_dict['media'] = [dict(m) for m in media]
        articles_list.append(article_dict)

    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    if settings is None:
        settings = {'logo_path': 'images/default-logo.png'}

    otp_enabled = session.get('otp_enabled', False)
    conn.close()

    
    return render_template('index.html', articles=articles_list, settings=settings, otp_enabled=otp_enabled)
import re

@app.route('/search')
@login_required
def search():
    conn = get_db_connection()
    try:
        query = request.args.get('query', '').lower().strip()
        sort = request.args.get('sort', 'desc')
        product = request.args.get('product', '')
        subproduct = request.args.get('subproduct', '')

        logging.debug(f"Iniciando busca: query={query}, sort={sort}, product={product}, subproduct={subproduct}")

        results = []

        # Dividir o termo de busca em palavras-chave
        keywords = [keyword.strip() for keyword in query.split() if keyword.strip()]
        logging.debug(f"Palavras-chave extraídas: {keywords}")

        if not keywords:
            logging.debug("Nenhuma palavra-chave fornecida, retornando resultados vazios")
            return render_template('search_results.html', query=query, results=[], settings={'logo_path': 'images/default-logo.png'})

        # Função auxiliar para construir a consulta de palavras-chave
        def add_keywords_query(base_query, fields, params):
            if keywords:
                conditions = []
                for field in fields:
                    for keyword in keywords:
                        conditions.append(f"LOWER(COALESCE({field}, '')) LIKE ?")
                        params.append(f'%{keyword}%')
                if conditions:
                    base_query += ' AND (' + ' OR '.join(conditions) + ')'
            return base_query, params

        # Busca em artigos (FAQs)
        articles_query = 'SELECT articles.*, users.username AS created_by_username FROM articles LEFT JOIN users ON articles.created_by = users.id WHERE 1=1'
        articles_params = []
        articles_query, articles_params = add_keywords_query(articles_query, ['title', 'question', 'solution'], articles_params)
        if product:
            articles_query += ' AND product = ?'
            articles_params.append(product)
        if subproduct and product == 'Creo Parametric':
            articles_query += ' AND subproduct = ?'
            articles_params.append(subproduct)
        if sort == 'asc':
            articles_query += ' ORDER BY created_at ASC'
        else:
            articles_query += ' ORDER BY created_at DESC'

        logging.debug(f"Query para artigos: {articles_query}, params: {articles_params}")
        articles = conn.execute(articles_query, articles_params).fetchall()
        logging.debug(f"Artigos encontrados: {len(articles)}")
        for article in articles:
            article_dict = dict(article)
            media = conn.execute('SELECT * FROM media WHERE article_id = ?', (article['id'],)).fetchall()
            article_dict['media'] = media
            results.append({
                'type': 'FAQ',
                'title': article_dict['title'],
                'snippet': article_dict['question'][:100] + ('...' if len(article_dict['question']) > 100 else ''),
                'url': url_for('view_article', id=article_dict['id']),
                'id': article_dict['id']  # Adicionado para consistência
            })
            for m in media:
                if any(keyword in m['media_path'].lower() for keyword in keywords) and m['media_path'].endswith(('.mp4', '.mov', '.mkv', '.avi', '.wmv')):
                    results.append({
                        'type': 'Vídeo',
                        'title': f"Vídeo: {article_dict['title']}",
                        'snippet': f"Vídeo relacionado a {article_dict['question'][:50]}...",
                        'url': url_for('static', filename=m['media_path'])
                    })

        # Busca em tutoriais (vídeos)
        tutorials_query = 'SELECT tutorials.*, users.username AS created_by_username FROM tutorials LEFT JOIN users ON tutorials.created_by = users.id WHERE 1=1'
        tutorials_params = []
        tutorials_query, tutorials_params = add_keywords_query(tutorials_query, ['title', 'description'], tutorials_params)
        if product:
            tutorials_query += ' AND product = ?'
            tutorials_params.append(product)
        if subproduct and product == 'Creo Parametric':
            tutorials_query += ' AND subproduct = ?'
            tutorials_params.append(subproduct)
        if sort == 'asc':
            tutorials_query += ' ORDER BY created_at ASC'
        else:
            tutorials_query += ' ORDER BY created_at DESC'

        logging.debug(f"Query para tutoriais: {tutorials_query}, params: {tutorials_params}")
        tutorials = conn.execute(tutorials_query, tutorials_params).fetchall()
        logging.debug(f"Tutoriais encontrados: {len(tutorials)}")
        for tutorial in tutorials:
            tutorial_dict = dict(tutorial)
            if any(keyword in tutorial_dict['video_path'].lower() for keyword in keywords) and tutorial_dict['video_path']:
                results.append({
                    'type': 'Vídeo',
                    'title': f"Vídeo: {tutorial_dict['title']}",
                    'snippet': tutorial_dict['description'][:100] + ('...' if len(tutorial_dict['description']) > 100 else ''),
                    'url': url_for('static', filename=tutorial_dict['video_path']),
                    'id': tutorial_dict['id']  # Adicionado para consistência
                })

        # Busca em procedimentos
        procedures_query = 'SELECT procedures.*, users.username AS created_by_username FROM procedures LEFT JOIN users ON procedures.created_by = users.id WHERE 1=1'
        procedures_params = []
        procedures_query, procedures_params = add_keywords_query(procedures_query, ['title', 'preview_text'], procedures_params)
        if product:
            procedures_query += ' AND product = ?'
            procedures_params.append(product)
        if subproduct and product == 'Creo Parametric':
            procedures_query += ' AND subproduct = ?'
            procedures_params.append(subproduct)
        if sort == 'asc':
            procedures_query += ' ORDER BY created_at ASC'
        else:
            procedures_query += ' ORDER BY created_at DESC'

        logging.debug(f"Query para procedimentos: {procedures_query}, params: {procedures_params}")
        procedures = conn.execute(procedures_query, procedures_params).fetchall()
        logging.debug(f"Procedimentos encontrados: {len(procedures)}")
        for procedure in procedures:
            procedure_dict = dict(procedure)
            results.append({
                'type': 'Procedimento',
                'title': procedure_dict['title'],
                'snippet': (procedure_dict['preview_text'][:100] + ('...' if procedure_dict['preview_text'] and len(procedure_dict['preview_text']) > 100 else '') if procedure_dict['preview_text'] is not None else 'Sem preview'),
                'url': url_for('view_procedure', id=procedure_dict['id']),
                'file_path': procedure_dict['file_path'],
                'created_by': procedure_dict['created_by'],
                'id': procedure_dict['id']  # Adicionado para o botão Editar
            })

        # Carregar settings
        settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
        if settings is None:
            settings = {'logo_path': 'images/default-logo.png'}

        logging.debug(f"Total de resultados: {len(results)}")
        conn.close()
        return render_template('search_results.html', query=query, results=results, settings=settings)
    except Exception as e:
        logging.error(f"Erro na busca: {str(e)}")
        conn.close()
        return render_template('search_results.html', query=query, results=[], settings={'logo_path': 'images/default-logo.png'}, error="Erro ao realizar a busca")

from flask import request, redirect, url_for, flash, render_template
from flask_login import login_user, current_user
from werkzeug.security import check_password_hash

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        print("Usuário já autenticado, redirecionando para index")
        return redirect(url_for('index'))

    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        print(f"Tentativa de login - Usuário: {username}, Senha: {password}")
        conn = get_db_connection()
        user = conn.execute('SELECT * FROM users WHERE username = ?', (username,)).fetchone()
        conn.close()

        if user and check_password_hash(user['password_hash'], password):
            user_obj = User(user['id'], user['username'], user['email'], user['role'], user['otp_secret'])
            login_user(user_obj)
            print("Login bem-sucedido")
            # Verificar explicitamente se otp_secret é None ou string vazia
            session['otp_enabled'] = user['otp_secret'] is not None and user['otp_secret'] != ""
            print(f"OTP Enabled: {session['otp_enabled']}, OTP Secret: {user['otp_secret']}")
            next_page = request.args.get('next')
            return redirect(next_page or url_for('index'))
        print("Erro de login: usuário ou senha inválidos")
        flash('Usuário ou senha inválidos. Tente novamente.', 'error')

    # Carregar settings para o template
    conn = get_db_connection()
    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    if settings is None:
        settings = {'logo_path': 'images/default-logo.png'}
    conn.close()

    print("Renderizando login.html com settings")
    return render_template('login.html', settings=settings)

@app.route('/logout')
@login_required
def logout():
    get_flashed_messages(with_categories=True)
    logout_user()
    flash('Você saiu da sua conta.', 'success')
    return redirect(url_for('login'))

@app.route('/edit_user', methods=['GET', 'POST'])
@login_required
def edit_user():
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        current_password = request.form['current_password']
        new_password = request.form['new_password']

        conn = get_db_connection()
        user = conn.execute('SELECT * FROM users WHERE id = ?', (current_user.id,)).fetchone()

        # Verificar senha atual
        if not check_password_hash(user['password_hash'], current_password):
            conn.close()
            flash('Senha atual incorreta.', 'error')
            return redirect(url_for('edit_user'))

        # Validar e-mail
        if not email.endswith('@infoaxis.com.br'):
            conn.close()
            flash('Apenas emails com domínio @infoaxis.com.br são permitidos.', 'error')
            return redirect(url_for('edit_user'))

        # Verificar se o username ou e-mail já está em uso (ignorando o próprio usuário)
        existing_user = conn.execute(
            'SELECT * FROM users WHERE (username = ? OR email = ?) AND id != ?',
            (username, email, current_user.id)
        ).fetchone()
        if existing_user:
            conn.close()
            flash('Nome de usuário ou e-mail já está em uso.', 'error')
            return redirect(url_for('edit_user'))

        # Atualizar dados
        password_hash = user['password_hash']
        if new_password:
            password_hash = generate_password_hash(new_password)

        try:
            conn.execute(
                'UPDATE users SET username = ?, email = ?, password_hash = ? WHERE id = ?',
                (username, email, password_hash, current_user.id)
            )
            conn.commit()
            flash('Dados atualizados com sucesso! Faça login novamente.', 'success')
            logout_user()  # Deslogar para aplicar as alterações
            return redirect(url_for('login'))
        except sqlite3.IntegrityError:
            conn.rollback()
            flash('Erro ao atualizar os dados. Nome de usuário ou e-mail já em uso.', 'error')
        finally:
            conn.close()

    # Carregar settings para o template
    conn = get_db_connection()
    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    conn.close()
    return render_template('edit_user.html', settings=settings)

from flask import request, url_for
import time

@app.route('/my_content')
@login_required
def my_content():
    start_time = time.time()
    conn = get_db_connection()
    
    # Parâmetros de filtro
    search = request.args.get('search', '')
    sort = request.args.get('sort', 'desc')
    product = request.args.get('product', '')
    subproduct = request.args.get('subproduct', '')

    # Função auxiliar para consultar com filtros
    def query_with_filters(table, params):
        query = f'SELECT * FROM {table} WHERE created_by = ?'
        query_params = [current_user.id]
        if search:
            query += ' AND title LIKE ?'
            query_params.append(f'%{search}%')
        if product:
            query += ' AND product = ?'
            query_params.append(product)
        if subproduct and product == 'Creo Parametric':
            query += ' AND subproduct = ?'
            query_params.append(subproduct)
        if sort == 'asc':
            query += ' ORDER BY created_at ASC'
        else:
            query += ' ORDER BY created_at DESC'
        return conn.execute(query, query_params).fetchall()

    # Consultas otimizadas
    articles_start = time.time()
    articles = query_with_filters('articles', [current_user.id])
    print(f"Tempo para consultar articles: {time.time() - articles_start:.3f} segundos")

    tutorials_start = time.time()
    tutorials = query_with_filters('tutorials', [current_user.id])
    print(f"Tempo para consultar tutorials: {time.time() - tutorials_start:.3f} segundos")

    procedures_start = time.time()
    procedures = query_with_filters('procedures', [current_user.id])
    print(f"Tempo para consultar procedures: {time.time() - procedures_start:.3f} segundos")

    # Carregar settings
    settings_start = time.time()
    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    if settings is None:
        settings = {'logo_path': 'images/default-logo.png'}
    print(f"Tempo para consultar settings: {time.time() - settings_start:.3f} segundos")

    print(f"Tempo total até renderizar: {time.time() - start_time:.3f} segundos")
    conn.close()

    return render_template('my_content.html', 
                          articles=articles, 
                          tutorials=tutorials, 
                          procedures=procedures, 
                          settings=settings)

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        
        print(f"Dados recebidos do formulário: username={username}, email={email}, password=****")
        print(f"Arquivos recebidos: {request.files}")

        if not email.endswith('@infoaxis.com.br'):
            flash('Apenas emails com domínio @infoaxis.com.br são permitidos.', 'error')
            print("Erro: Email não pertence ao domínio @infoaxis.com.br")
        else:
            try:
                file_path = None
                if 'file' in request.files:
                    file = request.files['file']
                    print(f"Arquivo recebido: {file.filename if file else 'Nenhum arquivo'}")
                    if file and file.filename != '' and allowed_file(file.filename):
                        filename = secure_filename(file.filename)
                        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename).replace('\\', '/')
                        file.save(file_path)
                        print(f"Arquivo salvo em: {file_path}")
                        flash(f'Arquivo {filename} enviado com sucesso!', 'success')
                    elif file and file.filename != '':
                        flash('Tipo de arquivo não permitido.', 'error')
                        print(f"Erro: Tipo de arquivo não permitido - {file.filename}")
                    else:
                        print("Nenhum arquivo válido enviado.")
                else:
                    print("Nenhum campo 'file' encontrado no request.files")

                # Gerar segredo OTP para Google Authenticator
                otp_secret = pyotp.random_base32()
                with get_db_connection() as conn:  # Usa with para fechar a conexão automaticamente
                    cursor = conn.cursor()
                    cursor.execute("INSERT INTO users (username, email, password_hash, role, otp_secret) VALUES (?, ?, ?, ?, ?)", 
                                   (username, email, generate_password_hash(password), 'user', otp_secret))
                    conn.commit()
                print(f"Usuário {username} registrado com sucesso com OTP Secret: {otp_secret}")
                flash('Cadastro realizado com sucesso! Configure o Google Authenticator escaneando o QR code abaixo.', 'success')
                return redirect(url_for('setup_otp', username=username))
            except sqlite3.IntegrityError as e:
                flash('Erro: Nome de usuário ou email já está em uso.', 'error')
                print(f"Erro ao registrar usuário: {e}")

    conn = get_db_connection()
    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    conn.close()  # Fecha a conexão manualmente aqui
    return render_template('register.html', settings=settings)

@app.route('/setup_otp/<username>', methods=['GET'])
def setup_otp(username):
    conn = get_db_connection()
    user = conn.execute('SELECT * FROM users WHERE username = ?', (username,)).fetchone()
    if not user:
        conn.close()
        flash('Usuário não encontrado.', 'error')
        return redirect(url_for('login'))

    # Gerar otp_secret se não existir
    if not user['otp_secret']:
        otp_secret = pyotp.random_base32()
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("UPDATE users SET otp_secret = ? WHERE username = ?", (otp_secret, username))
            conn.commit()
    else:
        otp_secret = user['otp_secret']

    otp_uri = pyotp.totp.TOTP(otp_secret).provisioning_uri(name=username, issuer_name='KnowledgeBase')
    qr = qrcode.make(otp_uri)
    img_io = BytesIO()
    qr.save(img_io, 'PNG')
    img_io.seek(0)

    img_base64 = base64.b64encode(img_io.getvalue()).decode('utf-8')
    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    if settings is None:
        settings = {'logo_path': 'images/default-logo.png'}
    conn.close()
    return render_template('setup_otp.html', username=username, otp_secret=otp_secret, qr_img=img_base64, settings=settings)

@app.route('/add', methods=['GET', 'POST'])
@login_required
def add_article():
    if request.method == 'POST':
        title = request.form['title']
        question = request.form['question']
        solution = request.form['solution']
        product = request.form['product']
        subproduct = request.form.get('subproduct', '')
        files = request.files.getlist('media')

        conn = get_db_connection()
        cursor = conn.cursor()
        created_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        cursor.execute('INSERT INTO articles (title, question, solution, created_at, created_by, product, subproduct) VALUES (?, ?, ?, ?, ?, ?, ?)', 
                       (title, question, solution, created_at, current_user.id, product, subproduct))
        article_id = cursor.lastrowid

        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename).replace('\\', '/')
                file.save(file_path)
                print(f"[INFO] Arquivo salvo: {file_path}")

                thumbnail_path = None
                thumbnail_filename = None
                if filename.lower().endswith(('.mp4', '.mov', '.mkv', '.avi', '.wmv')):
                    thumbnail_filename = f"{os.path.splitext(filename)[0]}_thumb.jpg"
                    thumbnail_path = os.path.join(app.config['THUMBNAIL_FOLDER'], thumbnail_filename).replace('\\', '/')
                    print(f"[INFO] Tentando gerar thumbnail para {filename}...")
                    if not create_thumbnail(file_path, thumbnail_path):
                        print(f"[ERROR] Falha ao gerar thumbnail para {filename}")
                        thumbnail_path = None
                    else:
                        print(f"[SUCCESS] Thumbnail gerado: {thumbnail_path}")

                cursor.execute('INSERT INTO media (article_id, media_path, thumbnail_path) VALUES (?, ?, ?)',
                               (article_id, f'uploads/{filename}', f'thumbnails/{thumbnail_filename}' if thumbnail_path else None))

        conn.commit()
        conn.close()
        flash('Artigo adicionado com sucesso!', 'success')
        return redirect(url_for('index'))

    conn = get_db_connection()
    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    conn.close()
    return render_template('add_article.html', settings=settings)

import re

def make_clickable(text):
    lines = text.split('\n')
    processed_lines = []
    for line in lines:
        if not line.strip():
            processed_lines.append('<br>')
            continue
        url_pattern = r'(https?://[^\s<>"]+|www\.[^\s<>"]+)'
        def replace_url(match):
            url = match.group(0)
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
            return f'<a href="{url}" target="_blank" rel="noopener noreferrer" class="text-blue-500 underline">{url}</a>'
        processed_line = re.sub(url_pattern, replace_url, line)
        processed_lines.append(processed_line)
    return '<div style="line-height: 1.0;">' + ''.join(processed_lines) + '</div>'

@app.route('/article/<int:id>')
@login_required
def view_article(id):
    conn = get_db_connection()
    article = conn.execute('SELECT articles.*, users.username as created_by_username FROM articles LEFT JOIN users ON articles.created_by = users.id WHERE articles.id = ?', (id,)).fetchone()
    if article is None:
        conn.close()
        flash('Artigo não encontrado.', 'error')
        return redirect(url_for('index'))

    media = conn.execute('SELECT * FROM media WHERE article_id = ?', (id,)).fetchall()
    article = dict(article)
    article['media'] = media

    # Depuração
    print(f"[DEBUG] Texto bruto de solution: {article['solution']}")

    # Processar o campo solution para quebras de linha e links
    article['solution'] = make_clickable(article['solution'].replace('\n', '<br>'))

    # Depuração
    print(f"[DEBUG] Texto processado de solution: {article['solution']}")

    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    conn.close()
    return render_template('view_article.html', article=article, settings=settings)

import os
from werkzeug.utils import secure_filename

@app.route('/delete_media/<int:media_id>', methods=['POST'])
@login_required
def delete_media(media_id):
    conn = get_db_connection()
    media = conn.execute('SELECT * FROM media WHERE id = ?', (media_id,)).fetchone()
    if media:
        article = conn.execute('SELECT * FROM articles WHERE id = ?', (media['article_id'],)).fetchone()
        if current_user.id != article['created_by'] and not current_user.is_admin:
            conn.close()
            flash('Você não tem permissão para excluir esta mídia.', 'error')
            return redirect(url_for('edit_article', id=media['article_id']))
        
        # Deletar o arquivo do sistema
        if os.path.exists(media['media_path']):
            os.remove(media['media_path'])
        
        # Deletar do banco de dados
        conn.execute('DELETE FROM media WHERE id = ?', (media_id,))
        conn.commit()
        flash('Mídia excluída com sucesso!', 'success')
    conn.close()
    return redirect(url_for('edit_article', id=media['article_id']))

@app.route('/edit_article/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_article(id):
    conn = get_db_connection()
    try:
        article = conn.execute('SELECT * FROM articles WHERE id = ?', (id,)).fetchone()
        if not article:
            flash('Artigo não encontrado ou você não tem permissão.', 'error')
            return redirect(url_for('index'))

        if request.method == 'POST':
            title = request.form['title']
            question = request.form['question']
            solution = request.form['solution']
            product = request.form['product']
            subproduct = request.form.get('subproduct', '')
            files = request.files.getlist('media')

            cursor = conn.cursor()
            try:
                cursor.execute('UPDATE articles SET title = ?, question = ?, solution = ?, product = ?, subproduct = ?, modified_at = ? WHERE id = ?',
                              (title, question, solution, product, subproduct, datetime.now().strftime('%Y-%m-%d %H:%M:%S'), id))

                # Processar novas mídias
                for file in files:
                    if file and allowed_file(file.filename):
                        filename = secure_filename(file.filename)
                        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename).replace('\\', '/')
                        file.save(file_path)
                        print(f"[DEBUG] Nova mídia salva em: {file_path}")
                        thumbnail_path = None
                        if filename.lower().endswith(('.mp4', '.mov', '.mkv', '.avi', '.wmv')):
                            thumbnail_filename = f"{os.path.splitext(filename)[0]}_thumb.jpg"
                            thumbnail_path = os.path.join(app.config['THUMBNAIL_FOLDER'], thumbnail_filename).replace('\\', '/')
                            if not create_thumbnail(file_path, thumbnail_path):
                                thumbnail_path = None
                            else:
                                print(f"[DEBUG] Thumbnail gerada em: {thumbnail_path}")
                        cursor.execute('INSERT INTO media (article_id, media_path, thumbnail_path) VALUES (?, ?, ?)',
                                      (id, f'uploads/{filename}', f'thumbnails/{thumbnail_filename}' if thumbnail_path else None))
                        print(f"[DEBUG] Inserido no banco: article_id={id}, media_path={f'uploads/{filename}'}, thumbnail_path={f'thumbnails/{thumbnail_filename}' if thumbnail_path else None}")

                conn.commit()
                flash('Artigo atualizado com sucesso!', 'success')
                return redirect(url_for('view_article', id=id))
            except sqlite3.Error as e:
                conn.rollback()
                print(f"[ERROR] Erro ao atualizar artigo: {str(e)}")
                flash('Erro ao atualizar o artigo. Tente novamente.', 'error')
                return redirect(url_for('edit_article', id=id))
            finally:
                cursor.close()

        media = conn.execute('SELECT * FROM media WHERE article_id = ?', (id,)).fetchall()
        settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone() or {'logo_path': 'images/default-logo.png'}
        return render_template('edit_article.html', article=dict(article), media=media, settings=settings)
    except Exception as e:
        print(f"[ERROR] Erro inesperado: {str(e)}")
        flash('Ocorreu um erro inesperado. Tente novamente.', 'error')
        return redirect(url_for('index'))
    finally:
        conn.close()
        print(f"[INFO] Conexão com o banco de dados fechada")

@app.route('/delete_article/<int:id>', methods=['POST'])
@login_required
def delete_article(id):
    if not current_user.is_admin:
        flash('Você não tem permissão para excluir este artigo.', 'error')
        return redirect(url_for('index'))

    conn = get_db_connection()
    article = conn.execute('SELECT * FROM articles WHERE id = ?', (id,)).fetchone()
    if article is None:
        conn.close()
        flash('Artigo não encontrado.', 'error')
        return redirect(url_for('index'))

    if current_user.id != article['created_by'] and not current_user.is_admin:
        conn.close()
        flash('Você não tem permissão para excluir este artigo.', 'error')
        return redirect(url_for('index'))

    media = conn.execute('SELECT media_path, thumbnail_path FROM media WHERE article_id = ?', (id,)).fetchall()
    for m in media:
        if m['media_path'] and os.path.exists(os.path.join('static', m['media_path'])):
            os.remove(os.path.join('static', m['media_path']))
        if m['thumbnail_path'] and os.path.exists(os.path.join('static', m['thumbnail_path'])):
            os.remove(os.path.join('static', m['thumbnail_path']))

    conn.execute('DELETE FROM media WHERE article_id = ?', (id,))
    conn.execute('DELETE FROM articles WHERE id = ?', (id,))
    conn.commit()
    conn.close()
    flash('Artigo excluído com sucesso!', 'success')
    return redirect(url_for('index'))

@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    if not current_user.is_admin:
        flash('Você não tem permissão para acessar as configurações.', 'error')
        return redirect(url_for('index'))

    conn = get_db_connection()
    if request.method == 'POST':
        file = request.files.get('logo')
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['IMAGE_FOLDER'], filename).replace('\\', '/')
            file.save(file_path)

            cursor = conn.cursor()
            cursor.execute('INSERT OR REPLACE INTO settings (id, logo_path) VALUES (1, ?)', (f'images/{filename}',))
            conn.commit()
            flash('Logo atualizado com sucesso!', 'success')
        else:
            flash('Por favor, selecione um arquivo de imagem válido (PNG, JPG, JPEG, GIF).', 'error')

    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    conn.close()
    return render_template('settings.html', settings=settings)

@app.route('/tutorials')
def tutorials():
    conn = get_db_connection()
    
    # Parâmetros de filtro
    search = request.args.get('search', '')
    sort = request.args.get('sort', 'desc')
    product = request.args.get('product', '')
    subproduct = request.args.get('subproduct', '')

    # Consulta com filtros
    query = 'SELECT tutorials.*, users.username as created_by_username, users.username as modified_by_username FROM tutorials LEFT JOIN users ON tutorials.created_by = users.id WHERE 1=1'
    query_params = []
    if search:
        query += ' AND (title LIKE ? OR description LIKE ?)'
        query_params.extend([f'%{search}%', f'%{search}%'])
    if product:
        query += ' AND product = ?'
        query_params.append(product)
    if subproduct and product == 'Creo Parametric':
        query += ' AND subproduct = ?'
        query_params.append(subproduct)
    if sort == 'asc':
        query += ' ORDER BY created_at ASC'
    else:
        query += ' ORDER BY created_at DESC'

    tutorials = conn.execute(query, query_params).fetchall()

    # Carregar settings
    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    if settings is None:
        settings = {'logo_path': 'images/default-logo.png'}

    conn.close()

    return render_template('tutorials.html', tutorials=tutorials, settings=settings)

@app.route('/procedures')
@login_required
def procedures():
    search_query = request.args.get('search', '')
    product_filter = request.args.get('product', '')
    subproduct_filter = request.args.get('subproduct', '')
    sort_order = request.args.get('sort', 'desc')

    conn = get_db_connection()
    query = '''
        SELECT p.*, u.username as created_by_username
        FROM procedures p
        JOIN users u ON p.created_by = u.id
        WHERE 1=1
    '''
    params = []

    if search_query:
        query += ' AND p.title LIKE ?'
        params.append(f'%{search_query}%')
    if product_filter:
        query += ' AND p.product = ?'
        params.append(product_filter)
    if subproduct_filter:
        query += ' AND p.subproduct = ?'
        params.append(subproduct_filter)

    query += ' ORDER BY p.created_at ' + ('DESC' if sort_order == 'desc' else 'ASC')

    procedures = conn.execute(query, params).fetchall()
    procedures_list = []

    for procedure in procedures:
        procedure_dict = dict(procedure)
        # Verificar e atribuir preview_path
        if procedure_dict.get('file_path', '').endswith('.pdf'):
            preview_filename = os.path.basename(procedure_dict['file_path']).replace('.pdf', '.png')
            preview_path = os.path.join('previews', preview_filename)
            full_preview_path = os.path.join('static', preview_path)
            if os.path.exists(full_preview_path):
                procedure_dict['preview_path'] = preview_path
            else:
                procedure_dict['preview_path'] = None
        else:
            procedure_dict['preview_path'] = procedure_dict.get('preview_path', None)
        procedures_list.append(procedure_dict)

    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    if settings is None:
        settings = {'logo_path': 'images/default-logo.png'}

    conn.close()
    return render_template('procedures.html', procedures=procedures_list, settings=settings)

@app.route('/add_procedure', methods=['GET', 'POST'])
@login_required
def add_procedure():
    if request.method == 'POST':
        try:
            title = request.form.get('title')
            product = request.form.get('product')
            subproduct = request.form.get('subproduct', '')
            file = request.files.get('file')

            if not title:
                return jsonify({"status": "error", "message": "Título é obrigatório"}), 400

            if not file or file.filename == '':
                return jsonify({"status": "error", "message": "Por favor, selecione um arquivo"}), 400

            if not allowed_file(file.filename, {'pdf', 'doc', 'docx'}):
                return jsonify({"status": "error", "message": "Tipo de arquivo não permitido. Use PDF, DOC ou DOCX"}), 400

            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename).replace('\\', '/')
            file.save(file_path)
            logging.debug(f"Arquivo salvo em: {file_path}")

            # Gerar preview para análise
            try:
                preview_path, preview_text = generate_preview(file_path, filename)
            except Exception as e:
                logging.error(f"Erro ao gerar preview: {str(e)}")
                return jsonify({"status": "error", "message": f"Erro ao gerar preview: {str(e)}"}), 500

            # Detectar produto automaticamente
            if not product or product not in ['Creo Parametric', 'Creo View', 'Windchill', 'Toolkit', 'Illustrate', 'Mathcad', 'Outros']:
                product = 'Outros'
                if 'creo' in title.lower() or (preview_text and 'creo' in preview_text.lower()):
                    product = 'Creo Parametric'
                elif 'windchill' in title.lower() or (preview_text and 'windchill' in preview_text.lower()):
                    product = 'Windchill'

            # Inserir no banco de dados
            conn = get_db_connection()
            cursor = conn.cursor()
            created_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            cursor.execute('INSERT INTO procedures (title, file_path, preview_path, preview_text, created_at, created_by, product, subproduct) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
                          (title, f'uploads/{filename}', preview_path, preview_text, created_at, current_user.id, product, subproduct))
            conn.commit()
            conn.close()

            return jsonify({
                "status": "success",
                "message": "Procedimento técnico adicionado com sucesso",
                "redirect": url_for('procedures')
            })
        except Exception as e:
            logging.error(f"Erro ao adicionar procedimento: {str(e)}")
            return jsonify({"status": "error", "message": str(e)}), 500

    # Para GET, renderizar o template
    conn = get_db_connection()
    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    conn.close()
    return render_template('add_procedure.html', settings=settings)

@app.route('/add_procedure_batch', methods=['GET', 'POST'])
@login_required
def add_procedure_batch():
    if request.method == 'POST':
        titles_input = request.form['titles'].strip()
        titles = [t.strip() for t in titles_input.split('\n') if t.strip()]
        pdf_files = request.files.getlist('pdf_files')
        word_files = request.files.getlist('word_files')
        product = request.form['product']
        subproduct = request.form.get('subproduct', '')

        if not pdf_files or not any(pdf_file.filename for pdf_file in pdf_files):
            flash('Por favor, selecione pelo menos um arquivo PDF.', 'error')
            return redirect(url_for('add_procedure_batch'))

        conn = get_db_connection()
        cursor = conn.cursor()

        for i, pdf_file in enumerate(pdf_files):
            if not pdf_file.filename:
                continue
            title = titles[i] if i < len(titles) else ''
            pdf_filename = secure_filename(pdf_file.filename)
            pdf_file_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename).replace('\\', '/')
            pdf_file.save(pdf_file_path)

            # Associar Word correspondente (se existir)
            word_file_path = None
            if i < len(word_files) and word_files[i].filename:
                word_filename = secure_filename(word_files[i].filename)
                word_file_path = os.path.join(app.config['UPLOAD_FOLDER'], word_filename).replace('\\', '/')
                word_files[i].save(word_file_path)
                print(f"Tentando salvar Word: {f'uploads/{word_filename}'}, procedure_id: {procedure_id}")  # Depuração
                cursor.execute('INSERT INTO media (article_id, media_path) VALUES (?, ?)',
                              (procedure_id, f'uploads/{word_filename}'))

            # Gerar preview para análise
            preview_path, preview_text = generate_preview(word_file_path if word_file_path else pdf_file_path,
                                                        word_filename if word_file_path else pdf_filename)

            # Se o título não for fornecido, usar o nome do arquivo sem extensão
            if not title:
                title = os.path.splitext(pdf_filename)[0]

            # Detectar produto automaticamente se não for fornecido ou for inválido
            if not product or product not in ['Creo Parametric', 'Windchill', 'Outros']:
                product = 'Outros'
                if 'creo' in title.lower() or (preview_text and 'creo' in preview_text.lower()):
                    product = 'Creo Parametric'
                elif 'windchill' in title.lower() or (preview_text and 'windchill' in preview_text.lower()):
                    product = 'Windchill'

            created_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            cursor.execute('INSERT INTO procedures (title, file_path, preview_path, preview_text, created_at, created_by, product, subproduct) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
                          (title, f'uploads/{pdf_filename}', preview_path, preview_text, created_at, current_user.id, product, subproduct))
            procedure_id = cursor.lastrowid

        conn.commit()
        conn.close()
        flash(f'{len(pdf_files)} procedimentos técnicos adicionados com sucesso!', 'success')
        return redirect(url_for('procedures'))

    conn = get_db_connection()
    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    conn.close()
    return render_template('add_procedure_batch.html', settings=settings)

from flask import Blueprint, render_template, redirect, url_for, request, flash
from flask_login import login_required, current_user
from datetime import datetime
from werkzeug.utils import secure_filename
import os

@app.route('/edit_procedure/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_procedure(id):
    conn = get_db_connection()
    try:
        procedure = conn.execute('SELECT * FROM procedures WHERE id = ?', (id,)).fetchone()
        if not procedure:
            conn.close()
            flash('Procedimento não encontrado.', 'error')
            return redirect(url_for('index'))

        # Carregar settings
        settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
        if settings is None:
            settings = {'logo_path': 'images/default-logo.png'}

        if request.method == 'POST':
            try:
                title = request.form.get('title')
                product = request.form.get('product')
                subproduct = request.form.get('subproduct', '')
                file = request.files.get('file')

                if not title:
                    flash('Título é obrigatório.', 'error')
                    return render_template('edit_procedure.html', procedure=procedure, settings=settings)

                file_path = procedure['file_path']
                if file and file.filename:
                    if not allowed_file(file.filename, {'pdf', 'doc', 'docx'}):
                        flash('Tipo de arquivo não permitido. Use PDF, DOC ou DOCX.', 'error')
                        return render_template('edit_procedure.html', procedure=procedure, settings=settings)
                    filename = secure_filename(file.filename)
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename).replace('\\', '/')
                    file.save(file_path)
                    logging.debug(f"Novo arquivo salvo em: {file_path}")

                # Atualizar no banco de dados
                conn.execute('UPDATE procedures SET title = ?, product = ?, subproduct = ?, file_path = ? WHERE id = ?',
                             (title, product, subproduct, file_path, id))
                conn.commit()
                flash('Procedimento atualizado com sucesso.', 'success')
                conn.close()
                return redirect(url_for('view_procedure', id=id))
            except Exception as e:
                logging.error(f"Erro ao editar procedimento: {str(e)}")
                flash(f'Erro ao atualizar procedimento: {str(e)}', 'error')
                return render_template('edit_procedure.html', procedure=procedure, settings=settings)

        conn.close()
        return render_template('edit_procedure.html', procedure=procedure, settings=settings)
    except Exception as e:
        logging.error(f"Erro na rota edit_procedure: {str(e)}")
        conn.close()
        flash(f'Erro ao acessar o procedimento: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/delete_procedure/<int:id>', methods=['POST'])
@login_required
def delete_procedure(id):
    if not current_user.is_admin:
        flash('Apenas administradores podem excluir procedimentos.', 'error')
        return redirect(url_for('procedures'))

    conn = get_db_connection()
    procedure = conn.execute('SELECT * FROM procedures WHERE id = ?', (id,)).fetchone()
    if procedure is None:
        conn.close()
        flash('Procedimento não encontrado.', 'error')
        return redirect(url_for('procedures'))

    if procedure['file_path'] and os.path.exists(os.path.join('static', procedure['file_path'])):
        os.remove(os.path.join('static', procedure['file_path']))
    if procedure['preview_path'] and os.path.exists(os.path.join('static', procedure['preview_path'])):
        os.remove(os.path.join('static', procedure['preview_path']))

    conn.execute('DELETE FROM procedures WHERE id = ?', (id,))
    conn.commit()
    conn.close()
    flash('Procedimento excluído com sucesso!', 'success')
    return redirect(url_for('procedures'))

@app.route('/edit_tutorial/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_tutorial(id):
    conn = get_db_connection()
    tutorial = conn.execute('SELECT * FROM tutorials WHERE id = ?', (id,)).fetchone()
    if tutorial is None:
        conn.close()
        flash('Tutorial não encontrado.', 'error')
        return redirect(url_for('tutorials'))

    if request.method == 'POST':
        title = request.form['title']
        description = request.form['description']
        product = request.form['product']
        subproduct = request.form.get('subproduct', '')
        video = request.files.get('video')

        video_path = tutorial['video_path']
        thumbnail_path = tutorial['thumbnail_path']
        thumbnail_filename = None

        if video and video.filename != '' and allowed_file(video.filename):
            if video_path and os.path.exists(os.path.join('static', video_path)):
                os.remove(os.path.join('static', video_path))
            if thumbnail_path and os.path.exists(os.path.join('static', thumbnail_path)):
                os.remove(os.path.join('static', thumbnail_path))

            filename = secure_filename(video.filename)
            video_full_path = os.path.join(app.config['VIDEO_FOLDER'], filename).replace('\\', '/')
            video.save(video_full_path)
            print(f"[INFO] Vídeo salvo: {video_full_path}")

            thumbnail_filename = f"{os.path.splitext(filename)[0]}_thumb.jpg"
            thumbnail_path = os.path.join(app.config['THUMBNAIL_FOLDER'], thumbnail_filename).replace('\\', '/')
            print(f"[INFO] Tentando gerar thumbnail para {filename}...")
            if not create_thumbnail(video_full_path, thumbnail_path):
                print(f"[ERROR] Falha ao gerar thumbnail para {filename}")
                thumbnail_path = None
            else:
                print(f"[SUCCESS] Thumbnail gerado: {thumbnail_path}")

            video_path = f'videos/{filename}'
            thumbnail_path = f'thumbnails/{thumbnail_filename}' if thumbnail_path else None

        cursor = conn.cursor()
        cursor.execute('''
            UPDATE tutorials 
            SET title = ?, description = ?, product = ?, subproduct = ?, 
                video_path = ?, thumbnail_path = ?, 
                modified_by_username = ?, modified_at = ? 
            WHERE id = ?
        ''', (title, description, product, subproduct, video_path, thumbnail_path, 
              current_user.username, datetime.now().strftime('%Y-%m-%d %H:%M:%S'), id))
        conn.commit()
        conn.close()
        flash('Tutorial atualizado com sucesso!', 'success')
        return redirect(url_for('tutorials'))

    # Fetch settings for GET request
    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    if settings is None:
        settings = {'logo_path': 'images/default-logo.png'}
    conn.close()
    return render_template('edit_tutorial.html', tutorial=tutorial, settings=settings)

@app.route('/delete_tutorial/<int:id>', methods=['POST'])
@login_required
def delete_tutorial(id):
    if not current_user.is_admin:
        flash('Apenas administradores podem excluir tutoriais.', 'error')
        return redirect(url_for('tutorials'))

    conn = get_db_connection()
    tutorial = conn.execute('SELECT * FROM tutorials WHERE id = ?', (id,)).fetchone()
    if tutorial is None:
        conn.close()
        flash('Tutorial não encontrado.', 'error')
        return redirect(url_for('tutorials'))

    if tutorial['video_path'] and os.path.exists(os.path.join('static', tutorial['video_path'])):
        os.remove(os.path.join('static', tutorial['video_path']))
    if tutorial['thumbnail_path'] and os.path.exists(os.path.join('static', tutorial['thumbnail_path'])):
        os.remove(os.path.join('static', tutorial['thumbnail_path']))

    conn.execute('DELETE FROM tutorials WHERE id = ?', (id,))
    conn.commit()
    conn.close()
    flash('Tutorial excluído com sucesso!', 'success')
    return redirect(url_for('tutorials'))

def create_thumbnail(video_path, thumbnail_path):
    print(f"[INFO] Iniciando create_thumbnail: video_path={video_path}, thumbnail_path={thumbnail_path}")
    try:
        abs_video_path = os.path.abspath(video_path)
        abs_thumbnail_path = os.path.abspath(thumbnail_path)
        print(f"[DEBUG {datetime.now().strftime('%H:%M:%S')}] Caminhos absolutos: video={abs_video_path}, thumbnail={abs_thumbnail_path}")

        if not os.path.exists(abs_video_path):
            print(f"[ERROR] Arquivo de vídeo não encontrado em {abs_video_path}")
            return False

        thumbnail_dir = os.path.dirname(abs_thumbnail_path)
        if not os.path.exists(thumbnail_dir):
            os.makedirs(thumbnail_dir, exist_ok=True)
            print(f"[INFO] Pasta {thumbnail_dir} criada.")
        if not os.access(thumbnail_dir, os.W_OK):
            print(f"[ERROR] Sem permissão de escrita na pasta {thumbnail_dir}")
            return False

        ffmpeg_cmd = app.config['FFMPEG_PATH']
        print(f"[INFO] Usando FFmpeg em: {ffmpeg_cmd}")

        if not os.path.exists(ffmpeg_cmd):
            print(f"[ERROR] FFmpeg não encontrado em {ffmpeg_cmd}. Ajuste app.config['FFMPEG_PATH'].")
            return False

        try:
            codecs_output = subprocess.run([ffmpeg_cmd, '-codecs'], capture_output=True, text=True, check=True)
            print(f"[DEBUG] Codecs suportados: {codecs_output.stdout}")
            if 'h264' not in codecs_output.stdout and 'h265' not in codecs_output.stdout and 'vp9' not in codecs_output.stdout:
                print("[ERROR] Codecs essenciais (h264, h265, vp9) não suportados pelo FFmpeg instalado.")
                return False
        except subprocess.CalledProcessError as e:
            print(f"[ERROR] Erro ao verificar codecs: {e.stderr}")
            return False

        try:
            ffmpeg_version = subprocess.run([ffmpeg_cmd, '-version'], capture_output=True, text=True, check=True)
            print(f"[INFO] FFmpeg encontrado: {ffmpeg_version.stdout.splitlines()[0]}")
        except subprocess.CalledProcessError as e:
            print(f"[ERROR] Erro ao executar FFmpeg: {e.stderr}")
            return False
        except Exception as e:
            print(f"[ERROR] Erro ao verificar FFmpeg: {str(e)}")
            return False

        print(f"[INFO] Executando comando FFmpeg...")
        process = subprocess.run([
            ffmpeg_cmd,
            '-i', abs_video_path,
            '-vf', 'scale=320:-1',
            '-frames:v', '1',
            '-q:v', '2',
            '-loglevel', 'debug',
            abs_thumbnail_path
        ], capture_output=True, text=True, check=False)

        if process.returncode == 0 and os.path.exists(abs_thumbnail_path):
            print(f"[SUCCESS] Thumbnail criado com sucesso em {abs_thumbnail_path}")
            return True
        else:
            print(f"[ERROR] Erro ao criar thumbnail (ffmpeg): Retorno {process.returncode}, Saída: {process.stderr}")
            return False
    except subprocess.CalledProcessError as e:
        print(f"[ERROR] Erro ao criar thumbnail (ffmpeg): {e.stderr}")
        return False
    except Exception as e:
        print(f"[ERROR] Erro inesperado ao criar thumbnail: {str(e)}")
        return False
    finally:
        print(f"[INFO] Finalizando create_thumbnail")

@app.route('/add_tutorial', methods=['GET', 'POST'])
@login_required
def add_tutorial():
    print(f"[INFO] Rota /add_tutorial acessada - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    conn = None
    try:
        # Estabelecer conexão com o banco de dados
        try:
            conn = get_db_connection()
            print(f"[INFO] Conexão com o banco de dados estabelecida")
        except Exception as e:
            print(f"[ERROR] Falha ao conectar ao banco de dados: {str(e)}")
            return jsonify({'status': 'error', 'message': f'Erro ao conectar ao banco de dados: {str(e)}'}), 500

        if request.method == 'POST':
            # Validar campos do formulário
            title = request.form.get('title')
            description = request.form.get('description')
            product = request.form.get('product')
            subproduct = request.form.get('subproduct', '') if product == 'Creo Parametric' else ''
            video = request.files.get('video')

            if not title or not description or not product:
                print(f"[ERROR] Campos obrigatórios ausentes: title={title}, description={description}, product={product}")
                return jsonify({'status': 'error', 'message': 'Campos obrigatórios (título, descrição, produto) não preenchidos.'}), 400

            if not video or not video.filename:
                print(f"[ERROR] Nenhum vídeo fornecido")
                return jsonify({'status': 'error', 'message': 'Nenhum vídeo foi fornecido.'}), 400

            if not allowed_file(video.filename):
                print(f"[ERROR] Tipo de arquivo não permitido: {video.filename}")
                return jsonify({'status': 'error', 'message': 'Tipo de arquivo não permitido. Use MP4, MOV, MKV, AVI ou WMV.'}), 400

            video_path = None
            converted_video_path = None
            thumbnail_path = None
            thumbnail_filename = None
            filename = secure_filename(video.filename)
            original_ext = filename.rsplit('.', 1)[1].lower()

            # Salvar o vídeo original temporariamente
            temp_video_path = os.path.join(app.config['VIDEO_FOLDER'], f"temp_{filename}").replace('\\', '/')
            print(f"[INFO] Tentando salvar vídeo temporário em: {temp_video_path}")
            try:
                os.makedirs(app.config['VIDEO_FOLDER'], exist_ok=True)
                if not os.access(app.config['VIDEO_FOLDER'], os.W_OK):
                    print(f"[ERROR] Sem permissão de escrita na pasta {app.config['VIDEO_FOLDER']}")
                    return jsonify({'status': 'error', 'message': 'Sem permissão para salvar o vídeo no servidor.'}), 500
                video.save(temp_video_path)
                print(f"[INFO] Vídeo temporário salvo com sucesso em: {temp_video_path}")
                if not os.path.exists(temp_video_path):
                    print(f"[ERROR] Vídeo temporário não encontrado em {temp_video_path} após salvamento")
                    return jsonify({'status': 'error', 'message': 'Erro: Arquivo de vídeo temporário não encontrado após salvamento.'}), 500
            except Exception as e:
                print(f"[ERROR] Falha ao salvar vídeo temporário: {str(e)}")
                return jsonify({'status': 'error', 'message': f'Falha ao salvar vídeo temporário: {str(e)}'}), 500

            # Converter para MP4 se for WMV
            if original_ext == 'wmv':
                converted_filename = f"{os.path.splitext(filename)[0]}.mp4"
                converted_video_path = os.path.join(app.config['VIDEO_FOLDER'], converted_filename).replace('\\', '/')
                print(f"[INFO] Convertendo {filename} para {converted_filename}...")
                try:
                    ffmpeg_cmd = app.config['FFMPEG_PATH']
                    subprocess.run([
                        ffmpeg_cmd,
                        '-i', temp_video_path,
                        '-c:v', 'libx264',
                        '-c:a', 'aac',
                        '-y',  # Sobrescrever sem perguntar
                        converted_video_path
                    ], check=True, capture_output=True, text=True)
                    print(f"[SUCCESS] Conversão concluída: {converted_video_path}")
                    # Remover o arquivo temporário .wmv
                    os.remove(temp_video_path)
                    video_path = f'videos/{converted_filename}'
                except subprocess.CalledProcessError as e:
                    print(f"[ERROR] Falha na conversão para MP4: {e.stderr}")
                    return jsonify({'status': 'error', 'message': f'Falha ao converter o vídeo para MP4: {str(e)}'}), 500
                except Exception as e:
                    print(f"[ERROR] Erro inesperado na conversão: {str(e)}")
                    return jsonify({'status': 'error', 'message': f'Erro inesperado ao converter o vídeo: {str(e)}'}), 500
            else:
                # Para outros formatos (mp4, mov, mkv, avi), mover o arquivo original
                video_path = os.path.join(app.config['VIDEO_FOLDER'], filename).replace('\\', '/')
                os.rename(temp_video_path, video_path)
                print(f"[INFO] Vídeo movido para: {video_path}")
                video_path = f'videos/{filename}'

            # Gerar thumbnail
            if filename.lower().endswith(('.mp4', '.mov', '.mkv', '.avi', '.wmv')):
                thumbnail_filename = f"{os.path.splitext(filename)[0]}_thumb.jpg"
                thumbnail_path = os.path.join(app.config['THUMBNAIL_FOLDER'], thumbnail_filename).replace('\\', '/')
                print(f"[INFO] Tentando gerar thumbnail para {filename}...")
                try:
                    os.makedirs(app.config['THUMBNAIL_FOLDER'], exist_ok=True)
                    if not os.access(app.config['THUMBNAIL_FOLDER'], os.W_OK):
                        print(f"[ERROR] Sem permissão de escrita na pasta {app.config['THUMBNAIL_FOLDER']}")
                        return jsonify({'status': 'error', 'message': 'Sem permissão para salvar o thumbnail no servidor.'}), 500
                    # Corrigir o caminho do vídeo para create_thumbnail
                    video_full_path = os.path.join('static/videos', filename).replace('\\', '/')
                    print(f"[DEBUG] Caminho ajustado do vídeo para create_thumbnail: {video_full_path}")
                    if create_thumbnail(video_full_path, thumbnail_path):
                        print(f"[SUCCESS] Thumbnail gerado: {thumbnail_path}")
                    else:
                        print(f"[WARNING] Falha ao gerar thumbnail para {video_full_path}. Prosseguindo sem thumbnail.")
                        thumbnail_path = None
                except Exception as e:
                    print(f"[ERROR] Erro ao gerar thumbnail: {str(e)}")
                    thumbnail_path = None  # Continuar mesmo se o thumbnail falhar

            # Salvar no banco de dados
            created_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            cursor = conn.cursor()
            try:
                cursor.execute(
                    'INSERT INTO tutorials (title, description, product, subproduct, video_path, thumbnail_path, created_at, created_by) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
                    (title, description, product, subproduct, video_path, f'thumbnails/{thumbnail_filename}' if thumbnail_path else None, created_at, current_user.id)
                )
                conn.commit()
                print(f"[INFO] Tutorial salvo no banco de dados com thumbnail_path: {f'thumbnails/{thumbnail_filename}' if thumbnail_path else 'None'}")
                return jsonify({'status': 'success', 'message': 'Vídeo tutorial adicionado com sucesso!', 'redirect': url_for('tutorials')})
            except sqlite3.IntegrityError as e:
                conn.rollback()
                print(f"[ERROR] Erro de integridade no banco de dados: {str(e)}")
                return jsonify({'status': 'error', 'message': f'Erro de integridade no banco de dados: {str(e)}'}), 500
            except Exception as e:
                conn.rollback()
                print(f"[ERROR] Falha ao salvar tutorial no banco de dados: {str(e)}")
                return jsonify({'status': 'error', 'message': f'Erro ao salvar tutorial no banco de dados: {str(e)}'}), 500

        # GET request
        settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
        return render_template('add_tutorial.html', settings=settings)

    except Exception as e:
        print(f"[ERROR] Erro inesperado no endpoint /add_tutorial: {str(e)}")
        return jsonify({'status': 'error', 'message': f'Erro inesperado no servidor: {str(e)}'}), 500

    finally:
        if conn is not None:
            conn.close()
            print(f"[INFO] Conexão com o banco de dados fechada")

@app.route('/view_tutorial/<int:id>')
@login_required
def view_tutorial(id):
    conn = get_db_connection()
    tutorial = conn.execute('SELECT tutorials.*, users.username as created_by_username FROM tutorials LEFT JOIN users ON tutorials.created_by = users.id WHERE tutorials.id = ?', (id,)).fetchone()
    if tutorial is None:
        conn.close()
        flash('Vídeo tutorial não encontrado.', 'error')
        return redirect(url_for('tutorials'))

    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    conn.close()
    return render_template('view_tutorial.html', tutorial=tutorial, settings=settings)

@app.route('/view_procedure/<int:id>')
@login_required
def view_procedure(id):
    conn = get_db_connection()
    procedure = conn.execute('SELECT procedures.*, users.username as created_by_username FROM procedures LEFT JOIN users ON procedures.created_by = users.id WHERE procedures.id = ?', (id,)).fetchone()
    if procedure is None:
        conn.close()
        flash('Procedimento não encontrado.', 'error')
        return redirect(url_for('procedures'))

    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    conn.close()
    return render_template('view_procedure.html', procedure=procedure, settings=settings)

@app.route('/download_procedure_word/<int:id>')
@login_required
def download_procedure_word(id):
    conn = get_db_connection()
    procedure = conn.execute('SELECT * FROM procedures WHERE id = ?', (id,)).fetchone()
    conn.close()

    if procedure is None:
        flash('Procedimento não encontrado.', 'error')
        return redirect(url_for('procedures'))

    # Criar documento Word
    doc = Document()
    doc.add_heading(procedure['title'], 0)
    doc.add_paragraph(f"Data de Criação: {procedure['created_at']}")
    if procedure['preview_text']:
        doc.add_paragraph(procedure['preview_text'])
    else:
        doc.add_paragraph("Nenhum texto de preview disponível.")

    # Salvar em buffer de memória
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Limpar caracteres inválidos do título para o nome do arquivo
    safe_title = ''.join(c for c in procedure['title'] if c.isalnum() or c in '-_').rstrip()
    download_name = f"{safe_title}.docx" if safe_title else f"procedure_{id}.docx"

    # Enviar o arquivo para download
    return send_file(
        buffer,
        as_attachment=True,
        download_name=download_name,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/download_original_word/<int:id>')
@login_required
def download_original_word(id):
    conn = get_db_connection()
    procedure = conn.execute('SELECT * FROM procedures WHERE id = ?', (id,)).fetchone()
    if procedure is None:
        conn.close()
        flash('Procedimento não encontrado.', 'error')
        return redirect(url_for('procedures'))

    # Buscar o arquivo Word associado na tabela media
    media = conn.execute('SELECT media_path FROM media WHERE article_id = ?', (id,)).fetchone()
    conn.close()

    if not media or not media['media_path'] or not media['media_path'].lower().endswith(('.doc', '.docx')):
        flash('Nenhum arquivo Word original encontrado para este procedimento.', 'error')
        return redirect(url_for('view_procedure', id=id))

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], media['media_path'].replace('uploads/', '')).replace('\\', '/')
    if not os.path.exists(file_path):
        flash('Arquivo Word original não encontrado no servidor.', 'error')
        return redirect(url_for('view_procedure', id=id))

    # Enviar o arquivo para download
    return send_from_directory(
        directory=os.path.dirname(file_path),
        path=os.path.basename(file_path),
        as_attachment=True,
        download_name=os.path.basename(file_path),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document' if file_path.lower().endswith('.docx') else 'application/msword'
    )

@app.route('/download_procedure_pdf/<int:id>')
@login_required
def download_procedure_pdf(id):
    conn = get_db_connection()
    procedure = conn.execute('SELECT * FROM procedures WHERE id = ?', (id,)).fetchone()
    conn.close()

    if procedure is None:
        flash('Procedimento não encontrado.', 'error')
        return redirect(url_for('procedures'))

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], procedure['file_path'].replace('uploads/', '')).replace('\\', '/')
    if not os.path.exists(file_path):
        flash('Arquivo PDF não encontrado no servidor.', 'error')
        return redirect(url_for('view_procedure', id=id))

    # Limpar caracteres inválidos do título para o nome do arquivo
    safe_title = ''.join(c for c in procedure['title'] if c.isalnum() or c in '-_').rstrip()
    download_name = f"{safe_title}.pdf" if safe_title else f"procedure_{id}.pdf"

    return send_from_directory(
        directory=os.path.dirname(file_path),
        path=os.path.basename(file_path),
        as_attachment=True,
        download_name=download_name,
        mimetype='application/pdf'
    )

import secrets  # Adicione isso no topo do arquivo se ainda não estiver

@app.route('/recover', methods=['GET', 'POST'])
def recover():
    if request.method == 'POST':
        username = request.form['username']
        otp_code = request.form['otp_code']
        
        conn = get_db_connection()
        user = conn.execute('SELECT * FROM users WHERE username = ?', (username,)).fetchone()
        conn.close()

        if user:
            totp = pyotp.TOTP(user['otp_secret'])
            if totp.verify(otp_code):
                # Gerar uma senha temporária aleatória
                temp_password = secrets.token_hex(8)  # Gera uma senha de 16 caracteres hexadecimais
                # Atualizar o password_hash no banco com o novo hash
                with get_db_connection() as conn:
                    cursor = conn.cursor()
                    cursor.execute("UPDATE users SET password_hash = ? WHERE username = ?", 
                                   (generate_password_hash(temp_password), username))
                    conn.commit()
                flash(f'Usuário: {user["username"]}. Use esta senha temporária para login: {temp_password} (redefina após login).', 'success')
                return redirect(url_for('login'))
            else:
                flash('Código OTP inválido ou expirado. Tente novamente.', 'error')
        else:
            flash('Usuário não encontrado.', 'error')

    conn = get_db_connection()
    settings = conn.execute('SELECT * FROM settings WHERE id = 1').fetchone()
    if settings is None:
        settings = {'logo_path': 'images/default-logo.png'}
    conn.close()
    return render_template('recover.html', settings=settings)


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5001)